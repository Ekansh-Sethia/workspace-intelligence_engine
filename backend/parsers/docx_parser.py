"""
Parser for Microsoft Word documents (.docx) using ``python-docx``.

Extracts paragraph text in reading order.
"""
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument

from parsers.base import BaseParser, Document
from utils.logger import logger


class DocxParser(BaseParser):
    """Handles .docx files."""

    def parse(self, file_content: bytes, filename: str, source_path: Optional[str] = None) -> Document:
        """
        Extract text from every paragraph of a DOCX file.

        Tables and embedded objects are intentionally skipped in V1 —
        only paragraph text is extracted.
        """
        import io
        doc = DocxDocument(io.BytesIO(file_content))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)
        rel = source_path or filename

        # Gather core properties if available
        meta: dict = {"parser": "DocxParser"}
        core = doc.core_properties
        if core.title:
            meta["title"] = core.title
        if core.author:
            meta["author"] = core.author

        if not full_text.strip():
            logger.warning(f"DocxParser: '{rel}' contains no extractable text")

        logger.info(f"DocxParser: parsed '{rel}' ({len(paragraphs)} paragraphs, {len(full_text)} chars)")

        return Document(
            filename=filename,
            source_path=rel,
            text=full_text,
            page_count=1,  # DOCX doesn't expose page count without rendering
            metadata=meta,
        )
